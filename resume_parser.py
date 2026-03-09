"""
parsers/resume_parser.py
─────────────────────────────────────────────────────────────────────────────
Handles loading and extracting plain text from resumes in three formats:
  • PDF   → pdfplumber (accurate multi-column extraction)
  • DOCX  → python-docx
  • TXT   → direct read

Public API
──────────
  parse_resume(file_path: str) -> str
      Returns the full extracted text of one resume.

  load_all_resumes(directory: str) -> list[dict]
      Returns a list of {"filename": str, "text": str} for every
      supported file in the given directory.
─────────────────────────────────────────────────────────────────────────────
"""

import os
import re
import logging

logger = logging.getLogger(__name__)

# ── Optional imports: gracefully skip if library not installed ────────────
try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    logger.warning("pdfplumber not installed — PDF parsing disabled. "
                   "Run: pip install pdfplumber")

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logger.warning("python-docx not installed — DOCX parsing disabled. "
                   "Run: pip install python-docx")

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


# ─────────────────────────────────────────────────────────────────────────
# Internal helpers
# ─────────────────────────────────────────────────────────────────────────

def _parse_pdf(file_path: str) -> str:
    """Extract text from a PDF using pdfplumber (handles multi-column layouts)."""
    if not PDF_SUPPORT:
        raise RuntimeError("pdfplumber is not installed. Run: pip install pdfplumber")

    pages_text = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)

    return "\n".join(pages_text)


def _parse_docx(file_path: str) -> str:
    """Extract text from a DOCX file, preserving paragraph order."""
    if not DOCX_SUPPORT:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also pull text from tables (skills are often in tables in DOCX resumes)
    table_cells = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    table_cells.append(cell.text.strip())

    return "\n".join(paragraphs + table_cells)


def _parse_txt(file_path: str) -> str:
    """Read a plain-text resume, trying UTF-8 then latin-1 as fallback."""
    for encoding in ("utf-8", "latin-1"):
        try:
            with open(file_path, "r", encoding=encoding) as fh:
                return fh.read()
        except UnicodeDecodeError:
            continue
    raise UnicodeDecodeError(f"Cannot decode {file_path} with utf-8 or latin-1")


def _clean_text(text: str) -> str:
    """
    Light normalisation: collapse excessive whitespace, remove null bytes,
    but keep newlines so date ranges (e.g. '2020 – 2023') stay intact.
    """
    text = text.replace("\x00", "")                 # null bytes
    text = re.sub(r"[ \t]{2,}", " ", text)          # multiple spaces/tabs → single space
    text = re.sub(r"\n{3,}", "\n\n", text)          # 3+ blank lines → 2
    return text.strip()


# ─────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────

def parse_resume(file_path: str) -> str:
    """
    Parse a single resume file and return its cleaned plain text.

    Parameters
    ----------
    file_path : str
        Absolute or relative path to the resume file.

    Returns
    -------
    str
        Extracted and cleaned resume text.

    Raises
    ------
    ValueError
        If the file extension is not supported.
    FileNotFoundError
        If the file does not exist.
    """
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"Resume not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: {SUPPORTED_EXTENSIONS}"
        )

    logger.info(f"Parsing resume: {file_path}")

    if ext == ".pdf":
        raw = _parse_pdf(file_path)
    elif ext == ".docx":
        raw = _parse_docx(file_path)
    else:  # .txt
        raw = _parse_txt(file_path)

    return _clean_text(raw)


def load_all_resumes(directory: str) -> list:
    """
    Scan a directory and parse every supported resume file.

    Parameters
    ----------
    directory : str
        Path to the folder containing resume files.

    Returns
    -------
    list of dict
        Each dict has:
          "filename" : original file name (no path)
          "filepath" : full path
          "text"     : extracted plain text
    """
    if not os.path.isdir(directory):
        raise NotADirectoryError(f"Resume directory not found: {directory}")

    results = []
    skipped = []

    for fname in sorted(os.listdir(directory)):
        ext = os.path.splitext(fname)[1].lower()
        if ext not in SUPPORTED_EXTENSIONS:
            continue

        fpath = os.path.join(directory, fname)
        try:
            text = parse_resume(fpath)
            results.append({
                "filename": fname,
                "filepath": fpath,
                "text":     text,
            })
            logger.info(f"  ✓ Parsed: {fname} ({len(text)} chars)")
        except Exception as exc:
            logger.warning(f"  ✗ Failed to parse {fname}: {exc}")
            skipped.append(fname)

    logger.info(
        f"Loaded {len(results)} resume(s) from '{directory}'. "
        f"Skipped: {len(skipped)}"
    )
    return results
